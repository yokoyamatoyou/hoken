import logging
from typing import List, Dict, Any
import pypdf
import docx
import requests
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

# Re-implement a simple Document class to avoid langchain_core dependency issues
class Document:
    def __init__(self, page_content: str, metadata: Dict[str, Any]):
        self.page_content = page_content
        self.metadata = metadata

    def __repr__(self):
        return f"Document(page_content='{self.page_content[:50]}...', metadata={self.metadata})"

def load_document(source: str) -> List[Document]:
    """
    Load documents from a file path or URL using basic libraries.
    This approach avoids heavy dependencies for stability.

    Args:
        source: The file path or URL to load.

    Returns:
        A list of Document objects.
    """
    docs = []
    try:
        if source.startswith(('http://', 'https://')):
            # It's a URL
            response = requests.get(source, headers={'User-Agent': 'Mozilla/5.0'})
            response.raise_for_status()  # Raise an exception for bad status codes
            soup = BeautifulSoup(response.content, 'html.parser')
            # Extract text from the body, removing script and style tags
            for script_or_style in soup(["script", "style"]):
                script_or_style.decompose()
            text = soup.body.get_text(separator='\\n', strip=True)
            metadata = {"source": source}
            docs.append(Document(page_content=text, metadata=metadata))

        elif source.lower().endswith('.pdf'):
            # It's a PDF file
            reader = pypdf.PdfReader(source)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    metadata = {"source": source, "page": i + 1}
                    docs.append(Document(page_content=text, metadata=metadata))

        elif source.lower().endswith('.docx'):
            # It's a Word document
            document = docx.Document(source)
            full_text = []
            for para in document.paragraphs:
                full_text.append(para.text)
            text = "\\n".join(full_text)
            metadata = {"source": source}
            docs.append(Document(page_content=text, metadata=metadata))

        else:
            raise ValueError(f"Unsupported source type: {source}")

        logger.info(f"Successfully loaded {len(docs)} document(s) from {source}")
        return docs

    except Exception as e:
        logger.error(f"Failed to load document from {source}: {e}", exc_info=True)
        return []
