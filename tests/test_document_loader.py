import logging
import os
import sys
from pathlib import Path

# Add project root to sys.path
project_root = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(project_root))

from src.document_loader import load_document

def test_load_document():
    # Setup logging
    logging.basicConfig(level=logging.INFO)

    # Test PDF
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter

        pdf_path = "dummy_test.pdf"
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.drawString(100, 750, "This is a test PDF document for the document loader.")
        c.drawString(100, 730, "It contains some sample text about insurance policies.")
        c.save()

        pdf_docs = load_document(pdf_path)
        assert pdf_docs is not None
        assert len(pdf_docs) > 0
        assert "insurance policies" in pdf_docs[0].page_content
        print(f"--- Loaded PDF ({len(pdf_docs)} docs) ---")
        print(pdf_docs[0].page_content[:200])
        print("-" * 20)
        os.remove(pdf_path)
    except ImportError:
        print("ReportLab not installed, skipping PDF test creation.")
    except Exception as e:
        print(f"Error during PDF test: {e}")
        assert False, f"PDF test failed: {e}"

    # Test DOCX
    try:
        import docx

        docx_path = "dummy_test.docx"
        doc = docx.Document()
        doc.add_paragraph("This is a test Word document.")
        doc.add_paragraph("It discusses the terms and conditions of an insurance plan.")
        doc.save(docx_path)

        docx_docs = load_document(docx_path)
        assert docx_docs is not None
        assert len(docx_docs) > 0
        assert "insurance plan" in docx_docs[0].page_content
        print(f"--- Loaded DOCX ({len(docx_docs)} docs) ---")
        print(docx_docs[0].page_content[:200])
        print("-" * 20)
        os.remove(docx_path)
    except ImportError:
        print("python-docx not installed, skipping DOCX test creation.")
    except Exception as e:
        print(f"Error during DOCX test: {e}")
        assert False, f"DOCX test failed: {e}"

    # Test URL
    # This requires an internet connection
    try:
        # A more reliable URL for testing
        url = "https://www.un.org/en/about-us/universal-declaration-of-human-rights"
        url_docs = load_document(url)
        assert url_docs is not None
        assert len(url_docs) > 0
        assert "Human Rights" in url_docs[0].page_content
        print(f"--- Loaded URL ({len(url_docs)} docs) ---")
        print(url_docs[0].page_content[:500].strip())
        print("-" * 20)
    except Exception as e:
        print(f"Error during URL test: {e}")
        assert False, f"URL test failed: {e}"

if __name__ == "__main__":
    test_load_document()
    print("All document loader tests passed!")
