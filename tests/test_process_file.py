import base64
from docx import Document
from PyPDF2 import PdfWriter
from PIL import Image
import openpyxl

from src.ui import main as GPT

ChatGPTClient = GPT.ChatGPTClient


def _client():
    return ChatGPTClient.__new__(ChatGPTClient)


def test_process_docx(tmp_path):
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Hello")
    doc.add_paragraph("World")
    doc.save(path)
    client = _client()
    result = client.process_file(str(path), ".docx")
    assert result == "Hello\nWorld"


def test_process_pdf_blank(tmp_path):
    path = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with open(path, "wb") as f:
        writer.write(f)
    client = _client()
    result = client.process_file(str(path), ".pdf")
    assert result == ""


def test_process_png(tmp_path):
    path = tmp_path / "img.png"
    img = Image.new("RGB", (1, 1), color="red")
    img.save(path)
    client = _client()
    result = client.process_file(str(path), ".png")
    with open(path, "rb") as f:
        expected = base64.b64encode(f.read()).decode("utf-8")
    assert result == expected


def test_process_xlsx(tmp_path):
    path = tmp_path / "sample.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append([1, 2])
    wb.save(path)
    client = _client()
    result = client.process_file(str(path), ".xlsx")
    assert "Excelファイル: 1個のシート" in result
    assert "【シート: Sheet1】" in result
    assert "行数: 1" in result
    assert "列数: 2" in result
    assert "行1: [1, 2]" in result
